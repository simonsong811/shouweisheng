"""Read-only inventory for the competition templates in the Documents folder."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


DOCUMENTS = Path(r"C:\Users\Owner\Documents")
OUT = Path(__file__).resolve().parents[1] / "tmp_template_inventory.json"


def docx_inventory(path: Path) -> dict:
    document = Document(path)
    return {
        "path": str(path),
        "paragraphs": [
            {"index": index, "text": paragraph.text}
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ],
        "tables": [
            {
                "index": table_index,
                "rows": [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ],
            }
            for table_index, table in enumerate(document.tables)
        ],
    }


def xlsx_inventory(path: Path) -> dict:
    workbook = load_workbook(path, data_only=False)
    return {
        "path": str(path),
        "sheets": [
            {
                "name": sheet.title,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "merged_ranges": [str(cell_range) for cell_range in sheet.merged_cells.ranges],
                "rows": [
                    [cell.value for cell in row]
                    for row in sheet.iter_rows()
                    if any(cell.value is not None for cell in row)
                ],
            }
            for sheet in workbook.worksheets
        ],
    }


def main() -> None:
    docx_files = []
    for path in DOCUMENTS.glob("*.docx"):
        document = Document(path)
        if len(document.paragraphs) <= 10 and len(document.tables) == 1:
            docx_files.append(path)
    xlsx_files = list(DOCUMENTS.glob("*.xlsx"))
    payload = {
        "docx": [docx_inventory(path) for path in docx_files],
        "xlsx": [xlsx_inventory(path) for path in xlsx_files],
    }
    OUT.write_text(json.dumps(payload, ensure_ascii=True, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
