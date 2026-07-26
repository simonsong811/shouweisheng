"""Build polished Word/PDF-ready submission documents and a clean export folder."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[1]
OUTPUT = WORKSPACE / "outputs" / "hand-hygiene-ai-submission"
EXPORT = OUTPUT / "final-export"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run_font(run, *, size=None, bold=None, color=None, east_asia="SimSun"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_section(section):
    # Named override from the Word preset baseline: A4 is standard for this Chinese submission.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.45)
    section.right_margin = Cm(2.45)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.05)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=GRAY)


def add_header_footer(document, header_text):
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = header_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=8.5, color=GRAY, east_asia="Microsoft YaHei")

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        add_page_number(paragraph)


def set_style_font(style, size, color=BLACK, bold=False, east_asia="SimSun"):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color


def configure_styles(document, preset):
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, BLACK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if preset == "proposal" else 6)
    normal.paragraph_format.line_spacing = 1.333 if preset == "proposal" else 1.25
    normal.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "proposal" else WD_ALIGN_PARAGRAPH.LEFT
    )
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12 if preset == "proposal" else 14, 6 if preset == "proposal" else 7),
        "Heading 3": (12, DARK_BLUE, 8 if preset == "proposal" else 10, 4 if preset == "proposal" else 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, size, color, True, "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, 11, BLACK, False)
        style.paragraph_format.left_indent = Cm(0.9525)
        style.paragraph_format.first_line_indent = Cm(-0.4775)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208 if preset == "proposal" else 1.25
        style.paragraph_format.widow_control = True


def create_decimal_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.append(tabs)
    paragraph_properties.append(indent)

    level.extend([start, number_format, level_text, level_justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = paragraph_properties.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        paragraph_properties.append(num_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(number)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=10, color=DARK_BLUE, east_asia="Microsoft YaHei")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_title_block(document, title, subtitle, kicker):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(kicker)
    set_run_font(run, size=10, bold=True, color=GRAY, east_asia="Microsoft YaHei")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, size=24, bold=True, color=BLACK, east_asia="Microsoft YaHei")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(14)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=12.5, color=BLUE, east_asia="Microsoft YaHei")


def markdown_to_docx(source, destination, *, preset, title, subtitle, kicker, header):
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document, preset)
    add_header_footer(document, header)
    add_title_block(document, title, subtitle, kicker)

    lines = source.read_text(encoding="utf-8").splitlines()
    skipped_title = False
    ordered_num_id = None
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            ordered_num_id = None
            continue
        if line.startswith("## "):
            ordered_num_id = None
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, line[3:])
            continue
        if line.startswith("### "):
            ordered_num_id = None
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, line[4:])
            continue
        if line.startswith("#### "):
            ordered_num_id = None
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline_runs(paragraph, line[5:])
            continue
        if line.startswith("- "):
            ordered_num_id = None
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, line[2:])
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            if ordered_num_id is None:
                ordered_num_id = create_decimal_numbering(document)
            paragraph = document.add_paragraph(style="List Number")
            apply_numbering(paragraph, ordered_num_id)
            add_inline_runs(paragraph, numbered.group(1))
            continue

        ordered_num_id = None
        paragraph = document.add_paragraph(style="Normal")
        if re.fullmatch(r"[^：:]{1,24}[：:]", line):
            paragraph.paragraph_format.keep_with_next = True
        add_inline_runs(paragraph, line)

    document.save(destination)


def copy_required_files():
    EXPORT.mkdir(parents=True, exist_ok=True)
    copies = {
        OUTPUT / "参赛作品信息表_手护智感_专家优化版.docx": EXPORT / "01_参赛作品信息表_手护智感_专家优化版.docx",
        OUTPUT / "参赛作品信息表_手护智感_专家优化版.pdf": EXPORT / "01_参赛作品信息表_手护智感_专家优化版.pdf",
        OUTPUT / "作者基本信息表_手护智感_待补充真实作者信息.xlsx": EXPORT / "02_作者基本信息表_手护智感_待补充真实作者信息.xlsx",
        WORKSPACE / "申报材料评审意见与修改说明.md": EXPORT / "05_申报材料评审意见与修改说明.md",
        WORKSPACE / "申报材料内容依据与文献核验.md": EXPORT / "06_申报材料内容依据与文献核验.md",
        WORKSPACE / "submission-checklist.md": EXPORT / "07_提交前检查清单.md",
    }
    for source, destination in copies.items():
        if not source.exists():
            raise FileNotFoundError(source)
        shutil.copy2(source, destination)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    proposal = OUTPUT / "手护智感_申报方案_优化版.docx"
    script = OUTPUT / "手护智感_演示视频脚本_优化版.docx"

    markdown_to_docx(
        WORKSPACE / "competition-application-optimized.md",
        proposal,
        preset="proposal",
        title="手护智感 AI 智能体",
        subtitle="基于 Wi-Fi CSI 与多源事件融合的非视觉化手卫生行为感知与提醒系统",
        kicker="第三届世界手卫生日 AI 设计大赛 · 申报方案",
        header="手护智感 AI 智能体 | 申报方案",
    )
    markdown_to_docx(
        WORKSPACE / "demo-video-script-optimized.md",
        script,
        preset="guide",
        title="手护智感 AI 智能体",
        subtitle="演示视频分镜、旁白与录制要求",
        kicker="第三届世界手卫生日 AI 设计大赛 · 演示脚本",
        header="手护智感 AI 智能体 | 演示视频脚本",
    )

    copy_required_files()
    shutil.copy2(proposal, EXPORT / "03_手护智感_申报方案_优化版.docx")
    shutil.copy2(script, EXPORT / "04_手护智感_演示视频脚本_优化版.docx")
    print(proposal)
    print(script)
    print(EXPORT)


if __name__ == "__main__":
    main()
