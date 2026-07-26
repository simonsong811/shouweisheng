"""Create a submission-ready copy of the competition information form.

Personal identity fields are intentionally left empty. They must be supplied by
the entrant before submission and cannot be inferred safely.
"""

from __future__ import annotations

from copy import copy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


WORKSPACE = Path(__file__).resolve().parents[1]
DOCUMENTS = Path(r"C:\Users\Owner\Documents")
OUTPUT = WORKSPACE / "outputs" / "hand-hygiene-ai-submission"

AGENT_NAME = "手护智感 AI 智能体"
FUNCTIONS = "工作场所提醒☑   科普宣传□   培训教学☑   文化营造□\n智能交互与个性化干预☑   其他□"
USERS = "临床医护☑   医技☑   后勤□   行政☑   患者或家属□   公众□\n其他☑（院感专职人员、质量管理人员）"
INTRODUCTION = (
    "“手护智感”是一款面向治疗室入口、ICU 床旁手消点与缓冲间的非视觉化手卫生 AI 智能体。"
    "它以“感知—推理—行动—学习”闭环工作：接收 Wi-Fi CSI、手消剂按压/水流、区域停留和设备"
    "状态等多源事件，识别经过、停留、疑似揉搓与干扰，结合 WHO 推荐时长和场景规则输出“疑似"
    "有效、时长不足、未发现、待复核”及置信度；对时长不足或未发现事件即时低打扰提醒，对低置信度"
    "或多人场景转人工复核，并将复核结果用于阈值校准和岗位复训建议。管理端统计疑似有效事件、"
    "时长不足事件（授权关联后可统计去标识化人次）、待复核量与设备在线率，帮助院感人员发现重点"
    "时段和场景。系统不采集视频、音频或人脸，默认仅保留去标识化事件；定位为重点空间行为初筛和"
    "改进助手，不替代 WHO 五时刻判断、手部覆盖质量评价和最终依从性质控。"
)
TECHNICAL = (
    "原型已实现微信小程序/网页管理端、本地联调网关、设备在线自检、事件融合、统计看板与复核队列。"
    "拟接入支持 CSI 的采集端和按压/水流传感器，经时间同步、滤波和动作片段切分后送入 CNN-LSTM/"
    "轻量 Transformer，输出动作类别、置信度和持续时间；规则代理结合按压、停留、场景阈值和设备"
    "质量，完成证据分级、提醒节流与复核路由。边缘网关采用 Node.js REST API，前端采用微信小程序"
    "原生框架及 HTML/CSS/JavaScript；原始 CSI 优先本地处理、最小化留存，模型版本、阈值与复核"
    "结果可审计。真实部署须经伦理、信息安全和场景验证后启用。"
)
ACCESS = (
    "1. 基于现有平台开发□\n\n"
    "2. 本地部署版本☑\n"
    "部署说明：在项目目录启动 gateway-dev/server.js；通过静态服务打开 index.html（推荐端口 "
    "8765），或用微信开发者工具导入小程序目录。网关地址：http://127.0.0.1:8787；测试场景："
    "treatment-room-01。\n"
    "体验账号：无需；真实数据接入前由管理员配置场景编号、设备白名单和访问权限。"
)


def find_source() -> Path:
    for path in DOCUMENTS.glob("*.docx"):
        document = Document(path)
        if len(document.paragraphs) <= 10 and len(document.tables) == 1:
            return path
    raise FileNotFoundError("Competition information form was not found.")


def set_cell_text(cell, text: str, size: float, center: bool = False) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = Pt(size)


def main() -> None:
    if len(INTRODUCTION) > 500:
        raise ValueError("Introduction exceeds the 500-character competition limit.")

    OUTPUT.mkdir(parents=True, exist_ok=True)
    document = Document(find_source())
    table = document.tables[0]

    set_cell_text(table.rows[0].cells[2], AGENT_NAME, 10.5)
    set_cell_text(table.rows[1].cells[2], FUNCTIONS, 10.5)
    set_cell_text(table.rows[2].cells[2], USERS, 10.5)
    set_cell_text(table.rows[4].cells[0], INTRODUCTION, 9.0)
    set_cell_text(table.rows[5].cells[2], TECHNICAL, 8.5)
    set_cell_text(table.rows[6].cells[2], ACCESS, 8.5)

    output = OUTPUT / "参赛作品信息表_手护智感_专家优化版.docx"
    document.save(output)
    print(output)
    print(f"introduction_chars={len(INTRODUCTION)}")


if __name__ == "__main__":
    main()
